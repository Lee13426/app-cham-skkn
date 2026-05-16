import streamlit as st
import google.generativeai as genai
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import io

# Thiết lập giao diện
st.set_page_config(page_title="AI Chấm SKKN - Đa nền tảng Model", layout="wide")

def doc_noi_dung(file):
    text = ""
    if file.name.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif file.name.endswith(".pdf"):
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text

def tao_file_word(ten_skkn, linh_vuc, ten_tac_gia, don_vi, ten_giam_khao, chuc_vu, noi_dung_ai):
    doc = docx.Document()
    
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n").bold = True
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PHIẾU NHẬN XÉT, ĐÁNH GIÁ\nGIẢI PHÁP ĐỀ NGHỊ CÔNG NHẬN SÁNG KIẾN\n(Phiếu dành cho thành viên chuyên ngành)")
    run_title.bold = True
    run_title.font.size = Pt(14)
    
    doc.add_paragraph(f"Họ và tên thành viên nhận xét: {ten_giam_khao}")
    doc.add_paragraph(f"Chức vụ: {chuc_vu}")
    doc.add_paragraph(f"Đơn vị công tác: {don_vi}\n")
    
    doc.add_paragraph("NHẬN XÉT, ĐÁNH GIÁ GIẢI PHÁP").runs[0].bold = True
    doc.add_paragraph(f"1. Tên giải pháp: {ten_skkn}")
    doc.add_paragraph(f"2. Thuộc lĩnh vực: {linh_vuc}")
    doc.add_paragraph(f"3. Tên tác giả (đồng tác giả): {ten_tac_gia}")
    doc.add_paragraph(f"4. Đơn vị công tác: {don_vi}\n")
    
    doc.add_paragraph("5. Nhận xét và Chấm điểm (Do AI thực hiện dựa trên Mẫu 09):").runs[0].bold = True
    
    doc.add_paragraph(noi_dung_ai)
    
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sign = p_sign.add_run("\nTHÀNH VIÊN NHẬN XÉT\n(Họ, tên và chữ ký)")
    run_sign.bold = True

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

with st.sidebar:
    st.title("⚙️ Cấu hình Hệ thống")
    api_key = st.text_input("Nhập Gemini API Key:", type="password")
    
    st.markdown("---")
    # Menu chọn Model AI cập nhật đầy đủ các bản mới nhất
    loai_model_name = st.selectbox(
        "🧠 Chọn phiên bản AI (Model):",
        [
            "gemini-1.5-flash (Tốc độ cực nhanh - Khuyên dùng)",
            "gemini-1.5-pro (Phân tích chuyên sâu - Khắt khe)",
            "gemini-flash-latest (Bản Flash tự động cập nhật mới nhất)",
            "gemini-pro-latest (Bản Pro tự động cập nhật mới nhất)",
            "gemini-flash-lite-latest (Bản Lite siêu nhẹ, siêu tốc)",
            "gemini-3-flash-preview (Bản Flash thế hệ 3 thử nghiệm)",
            "gemini-3.1-pro-preview (Bản Pro thế hệ 3.1 siêu việt)"
        ]
    )
    
    st.info("Hệ thống tuân thủ thang điểm 100 theo Mẫu 09 và xuất form Mẫu 10.")

st.header("📝 Thông tin Định danh (Xuất file Word)")
col1, col2 = st.columns(2)
with col1:
    ten_giam_khao = st.text_input("Tên thành viên nhận xét (Giám khảo):", value="Giám khảo AI")
    chuc_vu = st.text_input("Chức vụ của giám khảo:", value="Chuyên gia AI")
    ten_tac_gia = st.text_input("Tên tác giả SKKN:")
with col2:
    ten_skkn = st.text_input("Tên giải pháp/SKKN:")
    linh_vuc = st.selectbox("Thuộc lĩnh vực:", ["Toán học", "Vật lí", "Hóa học", "Sinh học", "Ngữ văn", "Khác"])
    don_vi = st.text_input("Đơn vị công tác (Trường):")

st.header("📄 Dữ liệu đầu vào & Tính năng nâng cao")
uploaded_file = st.file_uploader("Tải lên file SKKN của bạn (Word/PDF)", type=["docx", "pdf"])

st.markdown("**🔍 Các công cụ rà soát:**")
col_a, col_b = st.columns(2)
with col_a:
    check_ai = st.checkbox("🤖 Đánh giá % văn bản do AI (ChatGPT/Gemini) tạo ra")
with col_b:
    check_plagiarism = st.checkbox("🕵️ Rà soát dấu hiệu sao chép, đạo văn")

if st.button("Bắt đầu chấm điểm"):
    if not api_key or not uploaded_file:
        st.error("Vui lòng nhập API Key và tải file lên!")
    else:
        with st.spinner("Hệ thống đang phân tích chuyên sâu. Vui lòng đợi..."):
            try:
                noi_dung = doc_noi_dung(uploaded_file)
                
                genai.configure(api_key=api_key)
                
                # Trích xuất chính xác tên model gửi cho Google (Lấy chữ đầu tiên trước dấu cách)
                model_name = loai_model_name.split(" ")[0]
                model = genai.GenerativeModel(model_name)
                
                prompt = f"""
                Bạn là thành viên hội đồng chuyên ngành chấm Sáng kiến kinh nghiệm. Hãy đọc nội dung SKKN sau và viết Phiếu nhận xét.
                
                I. CHẤM ĐIỂM THEO THANG 100 ĐIỂM:
                1. Tính mới, tính sáng tạo: Tối đa 30 điểm. (Mức độ: Hoàn toàn mới 27-30đ, Cải tiến khá 21-26đ, Cải tiến trung bình 16-20đ, Cải tiến ít 1-15đ, Không có 0đ).
                2. Khả năng áp dụng: Tối đa 30 điểm. (Mức độ: Toàn tỉnh/ngoài tỉnh 27-30đ, Trong ngành/địa bàn tỉnh 21-26đ, Trong đơn vị 16-20đ, Ít khả năng 1-15đ, Không có 0đ).
                3. Hiệu quả của giải pháp: Tối đa 40 điểm. (Mức độ: Cao 31-40đ, Khá 21-30đ, Trung bình 11-20đ, Ít 1-10đ, Không 0đ).
                """
                
                if check_ai or check_plagiarism:
                    prompt += "\nII. RÀ SOÁT TÍNH TRUNG THỰC (BẮT BUỘC TRẢ LỜI NẾU CÓ YÊU CẦU):\n"
                    if check_ai:
                        prompt += "- Đánh giá AI: Dự đoán tỷ lệ % văn bản này do Trí tuệ nhân tạo (AI) viết. Chỉ ra rõ lý do (ví dụ: văn phong quá máy móc, dùng câu từ sáo rỗng thường thấy của AI, cấu trúc quá hoàn hảo nhưng thiếu tính thực tế cá nhân).\n"
                    if check_plagiarism:
                        prompt += "- Đánh giá Đạo văn: Chỉ ra các dấu hiệu cắt ghép, chép nhặt (ví dụ: các đoạn văn lủng củng không đồng nhất phong cách, số liệu minh chứng vô lý, hoặc nội dung giống hệt các mẫu giáo án/SKKN đại trà trên mạng).\n"

                prompt += """
                \nIII. TỔNG HỢP VÀ KẾT LUẬN:
                - Tổng điểm: .../100
                - Điều kiện xét duyệt: Tổng điểm >= 70, Tính mới >= 20, Hiệu quả >= 25.
                - Kết luận: [Đạt / Không đạt] (Chỉ ghi Đạt nếu thỏa mãn CẢ 3 điều kiện trên).
                
                Tuyệt đối không dùng ký tự Markdown dư thừa như # hay *.
                
                Nội dung SKKN:
                """
                prompt += f"\n{noi_dung}"
                
                response = model.generate_content(prompt)
                
                st.session_state.ai_text = response.text
                st.session_state.word_bytes = tao_file_word(ten_skkn, linh_vuc, ten_tac_gia, don_vi, ten_giam_khao, chuc_vu, response.text)
                
                st.success(f"Đã phân tích xong bằng bộ não {model_name}!")
                
            except Exception as e:
                st.error(f"Có lỗi xảy ra: {e}")

if "ai_text" in st.session_state:
    st.markdown("### 📊 Phiếu nhận xét (Bản xem trước)")
    st.write(st.session_state.ai_text)
    
    st.download_button(
        label="📥 TẢI XUỐNG FILE WORD",
        data=st.session_state.word_bytes,
        file_name="Phieu_Danh_Gia_SKKN_HoanChinh.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )
